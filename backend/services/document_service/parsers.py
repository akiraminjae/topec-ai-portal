"""문서 포맷별 구조 분석 + 텍스트 추출.

지원 포맷: docx, xlsx, pdf (파이썬 라이브러리로 직접 파싱),
hwpx/hwp (hwp_parser_service — Node.js 라이브러리 kordoc 기반 — 로 위임).

hwp(3.x/5.x)·hwpx는 자체 구현 대신 kordoc(https://github.com/chrisryugj/kordoc, MIT)을
쓰는 별도 서비스(`backend/services/hwp_parser_service`)에 위임합니다. 한글 문서 포맷은
복잡도가 높아 검증된 전문 라이브러리를 쓰는 편이 자체 구현보다 훨씬 안정적입니다.
"""
import io
import os

import httpx
from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader

HWP_PARSER_URL = os.environ.get("HWP_PARSER_URL", "http://hwp-parser-service:8107")


def extract(raw: bytes, fmt: str, filename: str = "document") -> dict:
    """raw 바이트를 fmt(확장자)에 맞는 파서로 처리해
    {"paragraphs": int, "tables": int, "images": int, "text": str} 를 반환합니다."""
    fmt = fmt.lower().lstrip(".")
    if fmt == "docx":
        return _parse_docx(raw)
    if fmt == "xlsx":
        return _parse_xlsx(raw)
    if fmt == "pdf":
        return _parse_pdf(raw)
    if fmt in ("hwpx", "hwp"):
        return _parse_via_hwp_parser_service(raw, filename, fmt)
    raise ValueError(f"지원하지 않는 포맷입니다: {fmt!r} (docx/xlsx/pdf/hwpx/hwp만 지원)")


def _parse_docx(raw: bytes) -> dict:
    doc = DocxDocument(io.BytesIO(raw))
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "images": len(doc.inline_shapes),
        "text": "\n".join(lines),
    }


def _parse_xlsx(raw: bytes) -> dict:
    wb = load_workbook(io.BytesIO(raw), data_only=True)
    lines = []
    images = 0
    for ws in wb.worksheets:
        lines.append(f"[시트: {ws.title}]")
        for row in ws.iter_rows():
            values = [str(c.value) for c in row if c.value is not None]
            if values:
                lines.append("\t".join(values))
        images += len(getattr(ws, "_images", []))
    return {
        "paragraphs": 0,
        "tables": len(wb.sheetnames),
        "images": images,
        "text": "\n".join(lines),
    }


def _parse_pdf(raw: bytes) -> dict:
    reader = PdfReader(io.BytesIO(raw))
    pages_text = [page.extract_text() or "" for page in reader.pages]
    images = 0
    for page in reader.pages:
        try:
            images += len(page.images)
        except Exception:
            pass
    text = "\n\n".join(pages_text)
    paragraphs = sum(1 for line in text.splitlines() if line.strip())
    return {"paragraphs": paragraphs, "tables": 0, "images": images, "text": text}


def _parse_via_hwp_parser_service(raw: bytes, filename: str, fmt: str) -> dict:
    try:
        resp = httpx.post(
            f"{HWP_PARSER_URL}/parse",
            files={"file": (filename, raw)},
            timeout=60,
        )
    except httpx.HTTPError as e:
        raise ValueError(f"hwp_parser_service({HWP_PARSER_URL}) 연결 실패: {e}")

    if resp.status_code != 200:
        try:
            detail = resp.json().get("detail", resp.text)
        except Exception:
            detail = resp.text
        raise ValueError(f"{fmt} 파싱 실패: {detail}")

    data = resp.json()
    return {
        "paragraphs": data.get("paragraphs", 0),
        "tables": data.get("tables", 0),
        "images": data.get("images", 0),
        "text": data.get("text", ""),
    }
